from django.http import FileResponse, HttpResponse
from rest_framework.views import APIView
from docx import Document
import pandas as pd
import json
import os
# for pdf
from io import BytesIO
from django.template.loader import get_template
from xhtml2pdf import pisa

from rest_framework import generics
from .models import MedicalData
from .serializers import MedicalDataSerializer

def render_to_pdf(template_src, context_dict={}):#'report.html',data
    # new code by Raghav garg
    template = get_template(template_src)
    html  = template.render(context_dict)
    result = BytesIO()
    pdf = pisa.pisaDocument(BytesIO(html.encode("ISO-8859-1")), result)#, link_callback=fetch_resources)
    if not pdf.err:
        return HttpResponse(result.getvalue(), content_type='application/pdf')
    return None

def data_collect_in_excel_file(data):

    # new code by Raghav garg
    file_name = './data.xlsx'
    if os.path.isfile(file_name):
        df_old = pd.read_excel(file_name)
        df_new = pd.concat([df_old, pd.DataFrame([data])])
    else:
        df_new = pd.DataFrame([data])
    df_new.to_excel(file_name, index=False)


class WordFileView(APIView):
    def post(self, request, format=None):
        data = json.loads(request.body)
        # print(data)

        # generate docx from data

        # doc = Document()
        # doc.add_heading(data['uhid'], level=1)
        # doc.add_paragraph(data['FullName'])
        # response = HttpResponse(content_type='application/msword')
        # response['Content-Disposition'] = 'attachment; filename="document.docx"'
        # doc.save(response)
        # return response
        
        # generate pdf from report.html

        pdf = render_to_pdf('report.html', data)
        if pdf:
            # Set response headers for download
            response = HttpResponse(pdf, content_type='application/pdf')
            response['Content-Disposition'] = 'inline; filename="Report.pdf"'
            return response
        else:
            return HttpResponse("Error generating PDF", status=500)


class MedicalDataListCreateView(generics.ListCreateAPIView):
    queryset = MedicalData.objects.all()
    serializer_class = MedicalDataSerializer